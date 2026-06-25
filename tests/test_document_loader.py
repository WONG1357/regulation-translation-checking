from pathlib import Path

import pytest
from docx import Document

from src.pipeline import process_document
from src.document_loader import _manual_page_for_section, load_document
from src.schemas import ProcessingSettings


SAMPLE_DOCX = Path("/Users/yichun/Downloads/QM01-N QUALITY MANUAL 质量手册 -20260427.docx")


def test_docx_loader_classifies_toc_and_skips_headings_while_preserving_body_order(tmp_path):
    path = tmp_path / "manual.docx"
    doc = Document()
    toc = doc.add_table(rows=2, cols=3)
    toc.rows[0].cells[0].text = "章节Chapter"
    toc.rows[0].cells[1].text = "标题Titles"
    toc.rows[0].cells[2].text = "页码Page"
    toc.rows[1].cells[0].text = "0.2"
    toc.rows[1].cells[1].text = "目的Purpose"
    toc.rows[1].cells[2].text = "7"
    doc.add_heading("目的Purpose", level=1)
    doc.add_paragraph("本手册之目的是描述质量管理体系。")
    doc.add_paragraph("The purpose of this manual is to describe the QMS.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "文件控制"
    table.rows[0].cells[1].text = "Document Control"
    doc.save(path)

    metadata, blocks = load_document(path)

    assert metadata.file_type == "docx"
    toc_blocks = [block for block in blocks if block.content_class == "table_of_contents"]
    assert toc_blocks
    assert all(block.ignored for block in toc_blocks)
    assert any(
        block.text == "目的Purpose"
        and block.block_type == "heading"
        and block.is_preserved_heading
        for block in blocks
    )
    body_blocks = [block for block in blocks if not block.ignored]
    assert [block.text for block in body_blocks] == [
        "目的Purpose",
        "本手册之目的是描述质量管理体系。",
        "The purpose of this manual is to describe the QMS.",
        "文件控制 | Document Control",
    ]
    assert all(block.section == "0.2" for block in body_blocks)
    assert all(block.page == 7 for block in body_blocks)


def test_page_range_filter_is_applied_after_docx_page_assignment(tmp_path):
    path = tmp_path / "manual.docx"
    doc = Document()
    toc = doc.add_table(rows=3, cols=3)
    toc.rows[0].cells[0].text = "章节Chapter"
    toc.rows[0].cells[1].text = "标题Titles"
    toc.rows[0].cells[2].text = "页码Page"
    toc.rows[1].cells[0].text = "1.0"
    toc.rows[1].cells[1].text = "公司简介Company History"
    toc.rows[1].cells[2].text = "20"
    toc.rows[2].cells[0].text = "2.0"
    toc.rows[2].cells[1].text = "质量方针Quality Policy"
    toc.rows[2].cells[2].text = "42"
    doc.add_heading("公司简介Company History", level=1)
    doc.add_paragraph("第一页中文。")
    doc.add_heading("质量方针Quality Policy", level=1)
    doc.add_paragraph("范围内中文。")
    doc.save(path)

    result = process_document(
        path,
        ProcessingSettings(
            dry_run=True,
            page_start=42,
            page_end=42,
            review_regulations=False,
            review_terminology=False,
        ),
    )

    assert result.metadata.page_count == 1
    assert {block.page for block in result.blocks} == {42}
    assert [block.text for block in result.blocks if not block.ignored] == [
        "质量方针Quality Policy",
        "范围内中文。",
    ]


def test_normal_subheading_matching_later_toc_entry_does_not_hijack_page_or_section(
    tmp_path,
):
    path = tmp_path / "manual.docx"
    doc = Document()
    toc = doc.add_table(rows=3, cols=3)
    toc.rows[0].cells[0].text = "章节Chapter"
    toc.rows[0].cells[1].text = "标题Titles"
    toc.rows[0].cells[2].text = "页码Page"
    toc.rows[1].cells[0].text = "2.0"
    toc.rows[1].cells[1].text = "质量方针和目标Quality Policy and Objective"
    toc.rows[1].cells[2].text = "23"
    toc.rows[2].cells[0].text = "5.3"
    toc.rows[2].cells[1].text = "质量方针Quality Policy"
    toc.rows[2].cells[2].text = "42"

    doc.add_heading("质量方针和目标Quality Policy and Objective", level=1)
    doc.add_paragraph("质量方针Quality Policy")
    doc.add_paragraph(
        "医科创建集团致力于成为国际顶级供应商，提供高品质，价格合理的器械。"
    )
    doc.add_paragraph(
        "MediConcepts Group’s vision is to be the world’s leading supplier."
    )
    doc.add_heading("质量方针Quality Policy", level=2)
    doc.add_paragraph("后续质量方针章节正文。")
    doc.save(path)

    _, blocks = load_document(path)

    page_23_texts = [block.text for block in blocks if block.page == 23]
    assert "质量方针和目标Quality Policy and Objective" in page_23_texts
    assert "质量方针Quality Policy" in page_23_texts
    assert any("国际顶级供应商" in text for text in page_23_texts)
    assert any("world’s leading supplier" in text for text in page_23_texts)
    assert all(
        block.section == "2.0"
        for block in blocks
        if block.page == 23 and not block.ignored
    )

    later_blocks = [block for block in blocks if "后续质量方针章节正文" in block.text]
    assert later_blocks
    assert later_blocks[0].page == 42
    assert later_blocks[0].section == "5.3"


def test_manual_page_resolver_uses_nearest_toc_parent_and_offsets():
    section_pages = {"1.0": 20, "3.2": 28, "7.3": 56}
    offsets: dict[str, int] = {}

    assert _manual_page_for_section("1.1", 26, section_pages, offsets) == 20
    assert _manual_page_for_section("7.3.3", 20, section_pages, offsets) == 56
    assert _manual_page_for_section("3.2.1.1", 32, section_pages, offsets) == 28
    assert _manual_page_for_section("3.2.4.1", 33, section_pages, offsets) == 29


@pytest.mark.skipif(not SAMPLE_DOCX.exists(), reason="sample quality manual not available")
def test_sample_docx_page_23_maps_quality_policy_to_section_2():
    result = process_document(
        SAMPLE_DOCX,
        ProcessingSettings(
            dry_run=True,
            page_start=23,
            page_end=23,
            review_regulations=False,
            review_terminology=False,
        ),
    )

    texts = [block.text for block in result.blocks]
    assert any("质量方针和目标Quality Policy and Objective" in text for text in texts)
    assert any("医科创建集团致力于成为国际顶级供应商" in text for text in texts)
    assert any("MediConcepts Group’s vision" in text for text in texts)
    assert all(
        block.section == "2.0"
        for block in result.blocks
        if not block.ignored and block.content_class != "table_of_contents"
    )


@pytest.mark.skipif(not SAMPLE_DOCX.exists(), reason="sample quality manual not available")
def test_sample_docx_page_33_contains_visible_responsibility_sections_only():
    result = process_document(
        SAMPLE_DOCX,
        ProcessingSettings(
            dry_run=True,
            page_start=33,
            page_end=33,
            review_regulations=False,
            review_terminology=False,
        ),
    )

    texts = [block.text for block in result.blocks]
    sections = {block.section for block in result.blocks}
    assert any("协助纠正预防措施行动的跟进" in text for text in texts)
    assert any("行政人事主管 Admin. & Personnel Supervisor" in text for text in texts)
    assert any("负责人力资源的规划" in text for text in texts)
    assert any("3.2.9.6" == section for section in sections)
    assert any("临床专家 Clinical Specialist" in text for text in texts)
    assert any("法規符合性負責人Person responsible for regulatory compliance" in text for text in texts)
    assert "3.2.11.1" in sections
    assert not any((block.section or "").startswith("3.1.1") for block in result.blocks)

    represented = sum(
        1
        for block in result.blocks
        if block.ignored
        or any(block.block_id in pair.source_block_ids for pair in result.pairs)
        or any(block.block_id == unpaired.block_id for unpaired in result.unpaired_blocks)
    )
    assert represented == len(result.blocks)
