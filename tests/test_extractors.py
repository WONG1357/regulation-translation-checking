import io
import unittest

from docx import Document

from src.extractors import extract_docx


class ExtractorTests(unittest.TestCase):
    def test_docx_preserves_paragraph_table_header_and_footer_text(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Quality Manual Header"
        doc.add_paragraph("2.0 质量方针和目标")
        doc.add_paragraph("质量方针应得到实施。")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "中文表格内容"
        table.cell(0, 1).text = "English table content"
        doc.sections[0].footer.paragraphs[0].text = "Page footer"
        output = io.BytesIO()
        doc.save(output)

        result = extract_docx("sample.docx", output.getvalue())
        texts = [block.text for block in result.blocks]

        self.assertIn("质量方针应得到实施。", texts)
        self.assertIn("中文表格内容", texts)
        self.assertIn("English table content", texts)
        self.assertIn("Quality Manual Header", texts)
        self.assertIn("Page footer", texts)
        self.assertEqual(
            {block.extraction_source for block in result.blocks},
            {"paragraph", "table", "header", "footer"},
        )


if __name__ == "__main__":
    unittest.main()
