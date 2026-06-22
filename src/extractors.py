from __future__ import annotations

import io
import re
from collections import defaultdict
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .utils import (
    MAX_FILE_SIZE_MB,
    SUPPORTED_EXTENSIONS,
    DocumentResult,
    TextBlock,
    contains_english,
    contains_chinese,
    normalize_text,
)
from .section_detection import assign_sections, parse_section_heading


HEADING_STYLE_RE = re.compile(r"heading|标题", re.I)
DOC_NUMBER_RE = re.compile(r"\b(?:doc(?:ument)?\.?\s*(?:no|number)|文件编号)[:：\s]*([A-Z0-9_\-/\.]+)", re.I)
REVISION_RE = re.compile(r"\b(?:rev(?:ision)?\.?|版本|修订)[:：\s]*([A-Z0-9_\-/\.]+)", re.I)


def validate_upload(file_name: str, file_size: int | None = None) -> str:
    suffix = Path(file_name).suffix.lower().lstrip(".")
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")
    if file_size and file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise ValueError(f"File is larger than {MAX_FILE_SIZE_MB} MB")
    return suffix


def extract_uploaded_file(file_name: str, data: bytes) -> DocumentResult:
    file_type = validate_upload(file_name, len(data))
    if not data:
        raise ValueError("Uploaded file is empty")
    if file_type == "docx":
        return extract_docx(file_name, data)
    if file_type == "pdf":
        return extract_pdf(file_name, data)
    if file_type == "txt":
        return extract_txt(file_name, data)
    raise ValueError(f"Unsupported file type: {file_type}")


def extract_docx(file_name: str, data: bytes) -> DocumentResult:
    doc = Document(io.BytesIO(data))
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    order = 0
    paragraph_number = 0
    current_heading: str | None = None

    table_index = 0
    for item in _iter_docx_body(doc):
        if isinstance(item, Paragraph):
            text = normalize_text(item.text)
            if not text:
                continue
            style_name = item.style.name if item.style else ""
            is_heading = bool(HEADING_STYLE_RE.search(style_name)) or _looks_like_heading(text)
            if is_heading:
                current_heading = text
            paragraph_number += 1
            blocks.append(
                TextBlock(
                    block_id=f"{file_name}:p:{paragraph_number}",
                    file_name=file_name,
                    file_type="docx",
                    text=text,
                    block_type="heading" if is_heading else "paragraph",
                    order=order,
                    section_heading=current_heading,
                    paragraph_number=paragraph_number,
                    extraction_source="paragraph",
                )
            )
            order += 1
        elif isinstance(item, Table):
            for row_index, row in enumerate(item.rows):
                row_texts = [normalize_text(cell.text) for cell in row.cells]
                for cell_index, cell_text in enumerate(row_texts):
                    for part_index, part in enumerate(_split_logical_lines(cell_text)):
                        blocks.append(
                            TextBlock(
                                block_id=f"{file_name}:t:{table_index}:r:{row_index}:c:{cell_index}:p:{part_index}",
                                file_name=file_name,
                                file_type="docx",
                                text=part,
                                block_type="table_cell",
                                order=order,
                                section_heading=current_heading,
                                table_index=table_index,
                                row_index=row_index,
                                cell_index=cell_index,
                                extraction_source="table",
                            )
                        )
                        order += 1
            table_index += 1

    for source_name, sections in (("header", doc.sections), ("footer", doc.sections)):
        seen: set[str] = set()
        for section_index, section in enumerate(sections):
            container = section.header if source_name == "header" else section.footer
            for item_index, paragraph in enumerate(container.paragraphs):
                text = normalize_text(paragraph.text)
                key = f"{section_index}:{text}"
                if not text or key in seen:
                    continue
                seen.add(key)
                blocks.append(
                    TextBlock(
                        block_id=f"{file_name}:{source_name}:{section_index}:{item_index}",
                        file_name=file_name,
                        file_type="docx",
                        text=text,
                        block_type=source_name,
                        order=order,
                        extraction_source=source_name,
                        is_repeated_header_footer=True,
                    )
                )
                order += 1

    if not blocks:
        warnings.append("No readable text found in DOCX.")

    assign_sections(blocks)
    all_text = "\n".join(block.text for block in blocks)
    metadata = extract_key_metadata(all_text)
    metadata["page_count"] = None
    return DocumentResult(file_name, "docx", blocks, warnings, metadata)


def _iter_docx_body(document: Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extract_pdf(file_name: str, data: bytes) -> DocumentResult:
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    pdf = fitz.open(stream=data, filetype="pdf")
    order = 0
    current_heading: str | None = None
    visual_elements = 0

    for page_idx in range(pdf.page_count):
        page = pdf.load_page(page_idx)
        raw_page_blocks = page.get_text("dict").get("blocks", [])
        page_visuals = [item for item in raw_page_blocks if item.get("type") != 0]
        visual_elements += len(page_visuals)
        page_items = _extract_pdf_page_items(page)
        if not page_items:
            warnings.append(f"Page {page_idx + 1}: no extractable text found; scanned/image-based content may require OCR.")
            continue
        for line_idx, item in enumerate(page_items):
            text = item["text"]
            is_heading = _looks_like_heading(text)
            if is_heading:
                current_heading = text
            blocks.append(
                TextBlock(
                    block_id=f"{file_name}:pg:{page_idx + 1}:b:{line_idx}",
                    file_name=file_name,
                    file_type="pdf",
                    text=text,
                    block_type=item["block_type"] if item["block_type"].startswith("table") else ("heading" if is_heading else "paragraph"),
                    order=order,
                    page_number=page_idx + 1,
                    section_heading=current_heading,
                    paragraph_number=line_idx + 1,
                    table_index=item.get("table_index"),
                    row_index=item.get("row_index"),
                    cell_index=item.get("cell_index"),
                    bbox=item.get("bbox"),
                    bbox_normalized=_normalize_bbox(item.get("bbox"), page.rect.width, page.rect.height),
                    extraction_source="table" if item["block_type"].startswith("table") else "paragraph",
                )
            )
            order += 1
        for visual_idx, visual in enumerate(page_visuals):
            bbox = tuple(float(value) for value in visual.get("bbox", (0, 0, 0, 0)))
            blocks.append(
                TextBlock(
                    block_id=f"{file_name}:pg:{page_idx + 1}:visual:{visual_idx}",
                    file_name=file_name,
                    file_type="pdf",
                    text=f"[Non-text visual element {visual_idx + 1}]",
                    block_type="image_or_visual",
                    order=order,
                    page_number=page_idx + 1,
                    section_heading=current_heading,
                    bbox=bbox,
                    bbox_normalized=_normalize_bbox(bbox, page.rect.width, page.rect.height),
                    extraction_source="unknown",
                )
            )
            order += 1

    assign_sections(blocks)
    _mark_repeated_headers_and_footers(blocks, pdf.page_count)

    if not blocks:
        warnings.append("No readable text found in PDF. The file may be scanned or image-based and require OCR.")

    all_text = "\n".join(block.text for block in blocks)
    metadata = extract_key_metadata(all_text)
    metadata["page_count"] = pdf.page_count
    metadata["ignored_visual_elements"] = visual_elements
    return DocumentResult(file_name, "pdf", blocks, warnings, metadata)


def extract_txt(file_name: str, data: bytes) -> DocumentResult:
    warnings: list[str] = []
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("utf-16", errors="ignore")
        warnings.append("Text file was not UTF-8; decoded with fallback encoding.")

    blocks: list[TextBlock] = []
    current_heading: str | None = None
    for idx, para in enumerate(re.split(r"\n\s*\n|\r\n\s*\r\n", text)):
        clean = normalize_text(para)
        if not clean:
            continue
        if _looks_like_heading(clean):
            current_heading = clean
        blocks.append(
            TextBlock(
                block_id=f"{file_name}:txt:{idx}",
                file_name=file_name,
                file_type="txt",
                text=clean,
                block_type="heading" if _looks_like_heading(clean) else "paragraph",
                order=len(blocks),
                section_heading=current_heading,
                paragraph_number=idx + 1,
                extraction_source="paragraph",
            )
        )

    if not blocks:
        warnings.append("No readable text found in TXT file.")
    assign_sections(blocks)
    metadata = extract_key_metadata(text)
    metadata["page_count"] = None
    return DocumentResult(file_name, "txt", blocks, warnings, metadata)


def extract_key_metadata(text: str) -> dict[str, str | None]:
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    title = next((line for line in lines[:20] if contains_chinese(line) or contains_english(line)), None)
    doc_no_match = DOC_NUMBER_RE.search(text)
    rev_match = REVISION_RE.search(text)
    return {
        "title": title,
        "document_number": doc_no_match.group(1) if doc_no_match else None,
        "revision": rev_match.group(1) if rev_match else None,
    }


def _looks_like_heading(text: str) -> bool:
    if parse_section_heading(text):
        return True
    if len(text) > 120:
        return False
    if re.match(r"^\s*(?:\d{1,2}\.\d+(?:\.\d+){0,3}|[A-Z]\.?)\s+.+", text):
        return True
    if "/" in text and contains_chinese(text) and contains_english(text) and len(text) < 80:
        return True
    return False


def _merge_pdf_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    buffer = ""
    for line in lines:
        if not buffer:
            buffer = line
            continue
        if len(buffer) < 80 and not re.search(r"[。.;；:]$", buffer):
            buffer = f"{buffer} {line}"
        else:
            merged.append(buffer)
            buffer = line
    if buffer:
        merged.append(buffer)
    return merged


def _extract_pdf_page_items(page) -> list[dict]:
    """Return text in visual reading order while preserving detected table cells."""
    table_items: list[dict] = []
    table_boxes: list[tuple[float, float, float, float]] = []
    try:
        found = page.find_tables()
        for table_index, table in enumerate(found.tables):
            table_boxes.append(tuple(float(value) for value in table.bbox))
            rows = table.extract()
            row_count = max(len(rows), 1)
            for row_index, row in enumerate(rows):
                col_count = max(len(row), 1)
                for cell_index, value in enumerate(row):
                    parts = _split_logical_lines(value or "")
                    if not parts:
                        continue
                    x0, y0, x1, y1 = table_boxes[-1]
                    # PyMuPDF does not expose every reconstructed cell rectangle
                    # consistently, so retain a stable approximate coordinate.
                    cell_bbox = (
                        x0 + (x1 - x0) * cell_index / col_count,
                        y0 + (y1 - y0) * row_index / row_count,
                        x0 + (x1 - x0) * (cell_index + 1) / col_count,
                        y0 + (y1 - y0) * (row_index + 1) / row_count,
                    )
                    for part_index, text in enumerate(parts):
                        part_height = (cell_bbox[3] - cell_bbox[1]) / len(parts)
                        part_bbox = (
                            cell_bbox[0],
                            cell_bbox[1] + part_height * part_index,
                            cell_bbox[2],
                            cell_bbox[1] + part_height * (part_index + 1),
                        )
                        table_items.append(
                            {
                                "text": text,
                                "block_type": "table_cell",
                                "table_index": table_index,
                                "row_index": row_index,
                                "cell_index": cell_index,
                                "bbox": part_bbox,
                            }
                        )
    except Exception:
        # Some PDFs have malformed vector geometry. Text extraction should
        # remain available even when table reconstruction fails.
        table_items = []
        table_boxes = []

    text_items: list[dict] = []
    page_dict = page.get_text("dict", sort=True)
    for raw_block in page_dict.get("blocks", []):
        if raw_block.get("type") != 0:
            continue
        for line in raw_block.get("lines", []):
            spans = line.get("spans", [])
            text = normalize_text("".join(span.get("text", "") for span in spans))
            if not text:
                continue
            bbox = tuple(float(value) for value in line.get("bbox", raw_block.get("bbox", (0, 0, 0, 0))))
            if any(_bbox_center_inside(bbox, table_bbox) for table_bbox in table_boxes):
                continue
            text_items.append({"text": text, "block_type": "paragraph", "bbox": bbox})

    items = text_items + table_items
    items.sort(
        key=lambda item: (
            round(item["bbox"][1], 1),
            item.get("table_index", -1),
            item.get("row_index", -1),
            item.get("cell_index", -1),
            round(item["bbox"][0], 1),
        )
    )
    return items


def _split_logical_lines(text: str) -> list[str]:
    return [
        clean
        for line in re.split(r"[\r\n]+", text or "")
        if (clean := normalize_text(line))
    ]


def _bbox_center_inside(
    bbox: tuple[float, float, float, float],
    container: tuple[float, float, float, float],
) -> bool:
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    return container[0] <= cx <= container[2] and container[1] <= cy <= container[3]


def _normalize_bbox(
    bbox: tuple[float, float, float, float] | None,
    page_width: float,
    page_height: float,
) -> tuple[float, float, float, float] | None:
    if not bbox or not page_width or not page_height:
        return None
    return (
        bbox[0] / page_width,
        bbox[1] / page_height,
        bbox[2] / page_width,
        bbox[3] / page_height,
    )


def _mark_repeated_headers_and_footers(blocks: list[TextBlock], page_count: int) -> None:
    if page_count < 2:
        return
    occurrences: dict[str, list[TextBlock]] = defaultdict(list)
    for block in blocks:
        if not block.bbox or not block.page_number:
            continue
        key = re.sub(r"\b\d+\b", "#", normalize_text(block.text).lower())
        if len(key) >= 3:
            occurrences[key].append(block)
    threshold = max(2, int(page_count * 0.5 + 0.5))
    for repeated in occurrences.values():
        pages = {block.page_number for block in repeated}
        if len(pages) < threshold:
            continue
        y_positions = [
            block.bbox_normalized[1]
            for block in repeated
            if block.bbox_normalized
        ]
        if not y_positions:
            continue
        # Repetition plus a consistently extreme vertical position is a much
        # safer signal than repetition alone (e.g. recurring table labels).
        near_top = max(y_positions) < 0.12
        near_bottom = min(y_positions) > 0.88
        if near_top or near_bottom:
            for block in repeated:
                block.is_repeated_header_footer = True
