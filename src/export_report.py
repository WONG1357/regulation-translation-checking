from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document

from .utils import safe_sheet_name


def dataframe_from_records(records: list[dict]) -> pd.DataFrame:
    return pd.DataFrame(records) if records else pd.DataFrame()


def build_excel_report(sheets: dict[str, list[dict] | pd.DataFrame]) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        wrote = False
        for name, data in sheets.items():
            df = data if isinstance(data, pd.DataFrame) else dataframe_from_records(data)
            if df.empty:
                df = pd.DataFrame([{"status": "No records"}])
            df.to_excel(writer, sheet_name=safe_sheet_name(name), index=False)
            wrote = True
        if not wrote:
            pd.DataFrame([{"status": "No report data"}]).to_excel(writer, sheet_name="Report", index=False)
    return output.getvalue()


def build_word_report(
    translation_issues: list[dict],
    terminology_issues: list[dict],
    regulatory_refs: list[dict],
    reference_comparison: list[dict],
) -> bytes:
    doc = Document()
    doc.add_heading("Bilingual Regulatory Document Review Report", 0)
    _add_table_section(doc, "Translation Issues", translation_issues)
    _add_table_section(doc, "Terminology Consistency Issues", terminology_issues)
    _add_table_section(doc, "Regulatory References", regulatory_refs)
    _add_table_section(doc, "Reference Regulation Comparison", reference_comparison)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _add_table_section(doc: Document, title: str, records: list[dict]) -> None:
    doc.add_heading(title, level=1)
    if not records:
        doc.add_paragraph("No records.")
        return
    keys = list(records[0].keys())
    table = doc.add_table(rows=1, cols=len(keys))
    table.style = "Table Grid"
    for idx, key in enumerate(keys):
        table.rows[0].cells[idx].text = str(key)
    for record in records[:100]:
        row = table.add_row().cells
        for idx, key in enumerate(keys):
            row[idx].text = str(record.get(key, ""))
